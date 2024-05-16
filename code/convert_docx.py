import argparse
import json
from docx import Document

'''
{ 
  "file": "/path/to/docx_file/file.docx"    
  "docx_pages" : 
    [{"page":"1","content":"xxx"},
     {"page":"2","content":"yyy"}]]
}
'''
def extract_pages_from_docx(docx_path):
    text = ""
    list = []
    docx_reader = Document(docx_path)
    num_pages = len(docx_reader.paragraphs)
    page_num = 1

    for paragraph in docx_reader.paragraphs:
        print(f"***** {page_num} / {num_pages} ****")
        print(f"{paragraph.paragraph_format.page_break_before}")
        if (paragraph.paragraph_format.page_break_before):
           text = paragraph.text()
           print(f"*****\n {text} \n****\n")
           value = { "page": page_num, "content": text}
           list.append(value)
        page_num = page_num + 1
        
    docx_information = { "file": docx_path, "pages": list}

    return docx_information
    
def save_docx_information_in_json(json_file, docx_dict):
        
        with open(json_file,'w', encoding='utf-8') as json_file:
            json.dump(docx_dict,json_file)
        
        return


def main(args):
    if args.docx_file_input is None:
        print(f"Docx file information is missing. '/your/folder/docx_file.docx'")
    else:
        input_docx = args.docx_file_input

    if args.json_file_output is None:
        print(f"JSON output file information is missing. '/your/folder/docx_file.json'")
    else:
        output_json = args.json_file_output
    
    docx_information = extract_pages_from_docx(input_docx)
    #save_docx_information_in_json( output_json, docx_information)

    return

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-i", "--docx_file_input", type=str)
    parser.add_argument("-o", "--json_file_output", type=str)
    args = parser.parse_args()
    main(args)